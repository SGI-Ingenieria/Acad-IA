from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_across_runs(paragraph, replacements):
    runs = paragraph.runs
    original = ''.join(run.text for run in runs)
    updated = original
    for source, target in replacements.items():
        updated = updated.replace(source, target)
    if updated == original or not runs:
        return
    runs[0].text = updated
    for run in runs[1:]:
        run.text = ''


def fix(path: Path, replacements):
    document = Document(path)
    for paragraph in paragraphs(document):
        replace_across_runs(paragraph, replacements)
    document.save(path)


fix(
    ROOT / 'supabase/template-assets/sep-rvoe/vigente/anexo-1-plan-render.docx',
    {
        '{d.fines_de_aprendizaje_o_formacion:convCRLF}':
            '{d.fines_de_aprendizaje_o_formacion:convCRLF()}',
        '{d.nombre_y_cargo_persona_facultada_para_autorizar_el_plan_de_estudios:convCRLF()}':
            '{d.nombre_y_cargo_de_la_persona_facultada_para_autorizar_el_plan_de_estudios:convCRLF()}',
    },
)

fix(
    ROOT / 'supabase/template-assets/sep-rvoe/vigente/anexo-3-programa-asignatura.docx',
    {
        '{d. numero_ciclo:convCRLF()}': '{d.numero_ciclo:convCRLF()}',
        '{d.datos. actividades_de_aprendizaje_bajo_conduccion_de_un_academico:convCRLF()}':
            '{d.datos.actividades_de_aprendizaje_bajo_conduccion_de_un_academico:convCRLF()}',
        '{d.datos. fines_de_aprendizaje_o_formacion:convCRLF()}':
            '{d.datos.fines_de_aprendizaje_o_formacion:convCRLF()}',
    },
)
